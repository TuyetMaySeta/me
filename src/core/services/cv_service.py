import asyncio
import base64
import io
import logging
import random
import re
import string
from typing import Any, Dict, List
from urllib.parse import parse_qs, urlparse

import aiohttp
import pandas as pd
import pdfplumber
from docx import Document
from langchain_core.messages import HumanMessage

from src.common.exception.exceptions import (
    ValidationException,
)
from src.common.llm_schema.system_promt import system_msg
from src.repository.employee_repository import EmployeeRepository

logger = logging.getLogger(__name__)


class CVService:
    """CV service with business logic for CV operations"""

    def __init__(
        self,
        employee_repository: EmployeeRepository,
        llm_instances: List[Any],
    ):
        self.employee_repository = employee_repository
        self.llm_instances = llm_instances
        self.num_llm_workers = len(llm_instances)
        self.max_concurrent_downloads = 10

    def _generate_cv_id(self, length: int = 6) -> str:
        """Generate random CV ID"""
        return "".join(random.choices(string.ascii_uppercase + string.digits, k=length))

    # --- Method để scan cv from file docx, pdf, image ---

    async def extract_cv_from_pdf(self, pdf_bytes: bytes, llm=None) -> Dict[str, Any]:
        """Calls LLM to extract CV data from PDF bytes."""
        text_content = await self._pdf_to_text(pdf_bytes)
        logger.info(f"{text_content}")
        if not text_content:
            raise ValueError("Extracted text is empty.")

        clean_text = self.fix_spaced_text(text_content)
        messages = [system_msg, HumanMessage(content=clean_text)]
        if llm is None:
            llm = self.llm_instances[0]
        cv_schema = await llm.ainvoke(messages)
        return cv_schema.model_dump()

    async def extract_cv_from_docx(self, docx_bytes: bytes, llm=None) -> Dict[str, Any]:
        """Calls LLM to extract CV data from DOCX bytes."""
        logger.info("Converting DOCX to PDF for text extraction")
        text_content = await self._docx_to_text(docx_bytes)
        if not text_content:
            raise ValueError("Extracted text is empty.")

        clean_text = self.fix_spaced_text(text_content)
        messages = [system_msg, HumanMessage(content=clean_text)]
        if llm is None:
            llm = self.llm_instances[0]
        cv_schema = await llm.ainvoke(messages)
        return cv_schema.model_dump()

    async def extract_cv_from_image(
        self, image_bytes: bytes, mime_type: str | None, llm=None
    ) -> Dict[str, Any]:
        """Calls LLM to extract CV data from image bytes."""
        logger.info(f"Extracting CV from image: {mime_type}")
        if llm is None:
            llm = self.llm_instances[0]
        image_base64 = await self._image_to_base64(image_bytes)
        message = {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": "Trích xuất thông tin CV từ ảnh này:",
                },
                {
                    "type": "image",
                    "source_type": "base64",
                    "data": image_base64,
                    "mime_type": mime_type,
                },
            ],
        }
        cv_schema = await llm.ainvoke([system_msg, message])
        return cv_schema.model_dump()

    # --- Method để insert vào db từ excel---

    def fix_spaced_text(self, text: str) -> str:
        def merge(match):
            return match.group(0).replace(" ", "")

        text = re.sub(r"((?:[A-Za-zÀ-ỹ]\s){2,}[A-Za-zÀ-ỹ])", merge, text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"(\n\s*){2,}", "\n\n", text)
        return text.strip()

    async def _pdf_to_text(self, content: bytes) -> str:
        """Extracts text from PDF bytes in a separate thread."""
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            lambda: "\n".join(
                page.extract_text() or ""
                for page in pdfplumber.open(io.BytesIO(content)).pages
            ),
        )

    async def _docx_to_text(self, content: bytes) -> str:
        """Extracts text from DOCX bytes in a separate thread."""
        document = Document(io.BytesIO(content))
        full_text = []

        # Extract header text
        for section in document.sections:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)

        # Extract main document text with paragraph formatting
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # Extract tables - critical for CV data
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        # Extract footer text
        for section in document.sections:
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)

        # Join with double newlines to preserve structure
        return "\n\n".join(full_text)

    async def _image_to_base64(self, image_bytes: bytes) -> str:
        return base64.b64encode(image_bytes).decode("utf-8")

    async def _download_and_extract_cv(
        self, session: aiohttp.ClientSession, url: str
    ) -> tuple[str, str]:
        """Downloads a CV from a URL and extracts its text."""
        download_url = None
        file_type = "pdf"

        match_drive = re.search(r"/file/d/([^/]+)", url)
        match_docs = re.search(r"/document/d/([^/]+)", url)

        if match_drive:
            file_id = match_drive.group(1)
            parsed = urlparse(url)
            query_params = parse_qs(parsed.query)
            resourcekey = query_params.get("resourcekey", [None])[0]

            download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
            if resourcekey:
                download_url += f"&resourcekey={resourcekey}"
        elif match_docs:
            file_id = match_docs.group(1)
            parsed = urlparse(url)
            query_params = parse_qs(parsed.query)
            resourcekey = query_params.get("resourcekey", [None])[0]
            download_url = (
                f"https://docs.google.com/document/d/{file_id}/export?format=pdf"
            )
            if resourcekey:
                download_url += f"&resourcekey={resourcekey}"
        else:
            raise ValueError("Unsupported URL format")

        timeout = aiohttp.ClientTimeout(total=60)  # Tăng timeout
        async with session.get(download_url, timeout=timeout) as resp:
            resp.raise_for_status()
            content = await resp.read()
            ctype = resp.headers.get("Content-Type", "").lower()
            if "image" in ctype:
                file_type = ctype
                return await self._image_to_base64(content), file_type

            return await self._pdf_to_text(content), file_type

    async def _llm_processor_worker(
        self, llm, download_queue: asyncio.Queue, results_list: list
    ):
        """
        Consumer: Takes text from queue, processes with LLM, and saves to DB.
        Appends a result tuple (status, excel_id, error_message) to the results list.
        """
        while True:
            item = await download_queue.get()
            if item is None:
                break

            excel_id, url, text_content, file_type = item
            cv_schema = None
            try:
                if not text_content:
                    raise ValueError("Extracted text is empty.")

                if "image" in file_type:
                    message = {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Trích xuất thông tin CV từ ảnh này:",
                            },
                            {
                                "type": "image",
                                "source_type": "base64",
                                "data": text_content,
                                "mime_type": file_type,
                            },
                        ],
                    }
                    cv_schema = await llm.ainvoke([system_msg, message])
                else:
                    clean_text = self.fix_spaced_text(text_content)
                    messages = [system_msg, HumanMessage(content=clean_text)]
                    cv_schema = await llm.ainvoke(messages)

                # 1. Process with LLM
                cv_data = cv_schema.model_dump()

                # 2. Prepare data for repository
                cv_data["id"] = self._generate_cv_id()
                cv_data["id_seta"] = excel_id

                # 3. Save to DB (synchronous call in a thread)
                await asyncio.to_thread(
                    self.employee_repository.create_employee, cv_data
                )

                # 4. Append success result
                results_list.append(("success", excel_id, None))
                logger.info(f"Successfully processed and saved CV for ID: {excel_id}")

            except Exception as e:
                # 5. Append failure result
                error_msg = f"Failed to process CV for ID {excel_id} from {url}: {e}"
                results_list.append(("fail", excel_id, error_msg))
                logger.error(error_msg)
            finally:
                download_queue.task_done()

    async def _downloader_worker(
        self,
        session: aiohttp.ClientSession,
        record: dict,
        download_queue: asyncio.Queue,
        semaphore: asyncio.Semaphore,
    ):
        """
        Producer: Downloads a single CV, extracts text, and puts it on the queue.
        """
        url = record.get("url_column")
        excel_id = record.get("id_column")

        async with semaphore:
            try:
                text_content, file_type = await self._download_and_extract_cv(
                    session, url
                )
                await download_queue.put((excel_id, url, text_content, file_type))
            except Exception as e:
                # Nếu download fail thì đánh case fail luôn
                await download_queue.put((excel_id, url, None, None))
                logger.error(f"Download error for ID {excel_id} at {url}: {e}")

    async def process_excel_bytes_sync(
        self, excel_bytes: bytes, url_column: str, id_column: str
    ) -> Dict[str, Any]:
        """
        Orchestrator: Reads an Excel file and processes all CVs concurrently.
        Returns a summary of successful and failed operations.
        """
        try:
            df = pd.read_excel(io.BytesIO(excel_bytes))
            if id_column not in df.columns or url_column not in df.columns:
                raise ValueError(
                    f"Excel must contain '{id_column}' and '{url_column}' columns."
                )
        except Exception as e:
            raise ValidationException(f"Invalid Excel file: {e}", "INVALID_EXCEL")

        records = (
            df[[id_column, url_column]]
            .rename(columns={id_column: "id_column", url_column: "url_column"})
            .to_dict("records")
        )

        download_queue = asyncio.Queue()
        results_list = []
        download_semaphore = asyncio.Semaphore(self.max_concurrent_downloads)

        async with aiohttp.ClientSession() as session:
            # --- Khởi tạo các workers ---
            # 1. LLM Consumers (chạy trước để sẵn sàng nhận việc)
            llm_consumers = [
                asyncio.create_task(
                    self._llm_processor_worker(llm, download_queue, results_list)
                )
                for llm in self.llm_instances
            ]

            # 2. Download Producers
            download_producers = [
                asyncio.create_task(
                    self._downloader_worker(
                        session, record, download_queue, download_semaphore
                    )
                )
                for record in records
            ]

            # --- Chờ đợi và dọn dẹp ---
            # 1. Chờ tất cả các download producer hoàn thành
            await asyncio.gather(*download_producers)

            # 2. Chờ queue được xử lý hết
            await download_queue.join()

            # 3. Gửi tín hiệu kết thúc cho các consumer
            for _ in range(self.num_llm_workers):
                await download_queue.put(None)

            # 4. Chờ các consumer kết thúc hoàn toàn
            await asyncio.gather(*llm_consumers)

        # --- Tổng hợp kết quả cuối cùng ---
        final_results = {"success_count": 0, "fail_count": 0, "failed_items": []}
        for status, excel_id, error_message in results_list:
            if status == "success":
                final_results["success_count"] += 1
            else:
                final_results["fail_count"] += 1
                final_results["failed_items"].append(
                    {"id": excel_id, "error": error_message}
                )

        logger.info(
            f"Finished processing. Success: {final_results['success_count']}, "
            f"Fail: {final_results['fail_count']}"
        )
        return final_results
